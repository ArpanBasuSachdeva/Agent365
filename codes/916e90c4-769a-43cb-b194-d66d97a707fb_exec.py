```python
from docx import Document
from docx.shared import Inches

def add_paragraph_to_docx(filepath, paragraph_text):
    """Adds a paragraph to an existing .docx file.

    Args:
        filepath: The path to the .docx file.
        paragraph_text: The text of the paragraph to add.

    Raises:
        FileNotFoundError: If the file does not exist.
        Exception: If any other error occurs during file processing.  Provides informative error message.

    """
    try:
        document = Document(filepath)
        document.add_paragraph(paragraph_text)
        document.save(filepath)
    except FileNotFoundError:
        raise FileNotFoundError(f"Error: File not found at path: {filepath}")
    except Exception as e:
        raise Exception(f"An error occurred while processing the file: {e}")


#Example usage.  Remember to replace with your actual file path and desired paragraph.
file_path = "New Microsoft Word Document.docx"
new_paragraph = "Computers are electronic devices that can accept data, process it according to a set of instructions (a program), and output the results.  They are essential tools in modern society, used for everything from simple calculations to complex simulations.  The development of computers has revolutionized many aspects of life, from communication and entertainment to science and industry."

add_paragraph_to_docx(file_path, new_paragraph)

```
