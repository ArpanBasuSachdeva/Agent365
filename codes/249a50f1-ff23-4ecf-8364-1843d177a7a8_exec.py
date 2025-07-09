```python
# Before running, ensure you have the required library installed:
# pip install python-docx

import docx
import os

def add_paragraph_to_docx(file_path: str):
    """
    Opens a .docx file, adds a new paragraph about computers to the end,
    and saves the changes back to the original file (in-place).

    Args:
        file_path (str): The path to the .docx file to be modified.
    """
    try:
        # Check if the file exists before attempting to open it
        if not os.path.exists(file_path):
            print(f"Error: The file '{file_path}' was not found.")
            # If the file doesn't exist, create a new one.
            document = docx.Document()
            print(f"A new document has been created at '{file_path}'.")
        else:
            # Open the existing document
            document = docx.Document(file_path)

        # The text for the new paragraph
        paragraph_text = (
            "Computers have become an integral part of modern life, revolutionizing "
            "industries, communication, and entertainment. From the massive mainframes "
            "of the mid-20th century to the powerful smartphones in our pockets, their "
            "evolution has been marked by a relentless pursuit of speed, efficiency, "
            "and miniaturization. At their core, computers process data through a "
            "combination of hardware and software, enabling complex calculations, "
            "data storage, and the execution of a vast array of applications that "
            "shape our daily experiences."
        )

        # Add the paragraph to the end of the document
        document.add_paragraph(paragraph_text)

        # Save the document, overwriting the original file
        document.save(file_path)

        print(f"Successfully added a paragraph to '{file_path}'.")

    except Exception as e:
        print(f"An error occurred while processing the file: {e}")

if __name__ == '__main__':
    # --- Configuration ---
    # Define the path to your Word document.
    # This script assumes the .docx file is in the same directory.
    # If not, provide the full path, e.g., 'C:/Users/YourUser/Documents/MyFile.docx'
    target_file_path = 'New Microsoft Word Document.docx'
    # -------------------

    # Call the function to modify the document
    add_paragraph_to_docx(target_file_path)
```