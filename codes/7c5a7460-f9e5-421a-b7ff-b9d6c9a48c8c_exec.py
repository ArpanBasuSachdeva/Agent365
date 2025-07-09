```python
import docx
import os

def add_paragraph_to_docx(file_path: str):
    """
    Opens a .docx file, adds a new paragraph about computers, and saves it.

    This script requires the 'python-docx' library.
    To install it, run: pip install python-docx

    Args:
        file_path (str): The full path to the .docx file to be modified.
    """
    try:
        # Check if the file exists and is a .docx file
        if not os.path.exists(file_path):
            print(f"Error: The file '{file_path}' does not exist.")
            # Create a new document if it doesn't exist, as per the example filename
            if file_path.lower().endswith('.docx'):
                 document = docx.Document()
                 print(f"Creating a new document at '{file_path}'.")
            else:
                print("Error: The provided path is not for a .docx file.")
                return
        else:
            document = docx.Document(file_path)

        # The paragraph to be added
        computer_paragraph = (
            "Computers have revolutionized the modern world, becoming indispensable "
            "tools in virtually every field of human endeavor. From complex "
            "scientific research and global finance to daily communication and "
            "entertainment, their processing power and ability to handle vast "
            "amounts of data have accelerated progress and transformed society. "
            "The continuous evolution of hardware and software promises even "
            "more profound changes in the years to come."
        )

        # Add the paragraph to the document
        document.add_paragraph(computer_paragraph)

        # Save the document in-place
        document.save(file_path)

        print(f"Successfully added a paragraph to '{file_path}'")

    except Exception as e:
        print(f"An error occurred: {e}")


if __name__ == '__main__':
    # --- Configuration ---
    # Replace this with the actual path to your DOCX file.
    # If the file doesn't exist, the script will create it.
    target_file = 'New Microsoft Word Document.docx'
    # --- End Configuration ---

    # Get the full path to the document in the current script's directory
    # This makes the script portable
    script_dir = os.path.dirname(os.path.abspath(__file__))
    full_path = os.path.join(script_dir, target_file)

    # Call the function to modify the document
    add_paragraph_to_docx(full_path)
```